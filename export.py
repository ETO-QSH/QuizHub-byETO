import os
import sys
import time
import shutil

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.enum.dml import MSO_THEME_COLOR
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH


def save_questions_to_word(questions, first_name, second_name, user):
    document = Document()

    style = document.styles['Normal']
    font_style = style.font
    font_style.name = '宋体'
    font_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    sorted_questions = sorted(questions, key=lambda x: tuple(map(int, x["uid"].split("-"))))
    document.add_heading(second_name, 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    for question in sorted_questions:
        uid = question.get('uid', 'N/A')
        body = question.get('question', 'N/A')
        answer = question.get('answer', 'N/A')
        explain = question.get('explain', 'N/A')
        options = question.get('options', {'√': '正确', '×': '错误'})

        p = document.add_paragraph()

        run_uid = p.add_run(f'{uid}')
        run_uid.bold = True
        run_uid.font.size = Pt(12)
        run_uid.font.name = 'Calibri (西文标题)'
        run_uid.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1

        run_body = p.add_run(f' {body}')
        run_body.bold = True

        for option, content in options.items():
            option_paragraph = document.add_paragraph()
            if option == answer:
                option_run = option_paragraph.add_run(f'{option}.{content}')
                option_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            else:
                option_paragraph.add_run(f'{option}.{content}')

        explain_paragraph = document.add_paragraph("解析：")
        explain_run = explain_paragraph.add_run(explain)
        explain_run.font.color.rgb = RGBColor(0, 0, 255)

        document.add_paragraph()

    APP_DIR = Path(sys.executable).parent if getattr(sys, 'frozen', False) else Path(__file__).parent
    export_dir = APP_DIR / "export" / user
    shutil.rmtree(export_dir, ignore_errors=True)
    export_dir.mkdir(parents=True, exist_ok=True)
    save_time = time.strftime("%Y-%m-%d-%H-%M-%S", time.localtime())
    docx_name = f'{first_name} - {second_name} - {save_time}.docx'
    docx_path = os.path.join(export_dir, docx_name)
    document.save(docx_path)
    return docx_name, docx_path
